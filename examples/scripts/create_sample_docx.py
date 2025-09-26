#!/usr/bin/env python3
"""Create a sample DOCX file for testing."""

try:
    from docx import Document
    
    # Create a new document
    doc = Document()
    
    # Add title
    title = doc.add_heading('LEGAL MEMORANDUM', 0)
    
    # Add content
    doc.add_heading('RE: Contract Analysis - TechServices Agreement', level=1)
    
    p1 = doc.add_paragraph('Date: September 25, 2024\n')
    p1.add_run('To: Legal Department\n').bold = True
    p1.add_run('From: Contract Review Team\n').bold = True
    
    doc.add_heading('Executive Summary', level=2)
    
    doc.add_paragraph(
        'This memorandum provides analysis of the TechServices LLC agreement dated January 15, 2024. '
        'Key findings include standard service terms, reasonable financial obligations, and '
        'appropriate termination clauses.'
    )
    
    doc.add_heading('Key Contract Terms', level=2)
    
    # Add a table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Term Category'
    hdr_cells[1].text = 'Details'
    
    # Add rows
    terms = [
        ('Parties', 'Acme Corporation and TechServices LLC'),
        ('Services', 'Software development, consulting, system maintenance'),
        ('Financial Terms', '$25,000 monthly retainer, $150 hourly rate'),
        ('Duration', 'January 15, 2024 - June 30, 2024'),
        ('Termination', '30 days written notice required'),
        ('Governing Law', 'New York State')
    ]
    
    for term, detail in terms:
        row_cells = table.add_row().cells
        row_cells[0].text = term
        row_cells[1].text = detail
    
    doc.add_heading('Legal Analysis', level=2)
    
    doc.add_paragraph(
        'The agreement contains standard commercial terms typical for technology service contracts. '
        'Notable provisions include:'
    )
    
    # Add bullet points
    doc.add_paragraph('Reasonable compensation structure with retainer and hourly components', style='List Bullet')
    doc.add_paragraph('Appropriate confidentiality protections for both parties', style='List Bullet')
    doc.add_paragraph('Clear deliverable and timeline requirements', style='List Bullet')
    doc.add_paragraph('Standard termination clause with adequate notice period', style='List Bullet')
    
    doc.add_heading('Recommendations', level=2)
    
    doc.add_paragraph(
        'The contract appears commercially reasonable and legally sound. We recommend proceeding '
        'with execution subject to final review of insurance and indemnification clauses.'
    )
    
    # Save the document
    doc.save('sample_docs/legal_memo.docx')
    print("Created sample_docs/legal_memo.docx")
    
except ImportError:
    print("python-docx not available - skipping DOCX file creation")
except Exception as e:
    print(f"Error creating DOCX file: {e}")