"""
File processor that handles multiple document formats (.pdf, .docx, .doc, .txt).
Replaces the PDF-only extraction with multi-format support.
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Dict, Any, Optional, List

# Import text extraction libraries
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# For .doc files (legacy Word format), we'd need python-docx2txt or similar
try:
    import docx2txt
    DOC_AVAILABLE = True
except ImportError:
    DOC_AVAILABLE = False

import sys
from pathlib import Path

# Add parent directories to path to find modules
sys.path.insert(0, str(Path(__file__).parent.parent))
from common_simple import console


class FileProcessor:
    """Handles text extraction from multiple document formats."""
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        self.config = config or {}
        self.ocr_enabled = self.config.get("ocr_enabled", True)
        self.clean_text = self.config.get("clean_text", True)
        
        # Check which formats are supported
        self.supported_formats = self._get_supported_formats()
    
    def _get_supported_formats(self) -> List[str]:
        """Get list of supported file formats based on available libraries."""
        formats = [".txt"]  # Always supported
        
        if PYMUPDF_AVAILABLE:
            formats.append(".pdf")
        
        if DOCX_AVAILABLE:
            formats.append(".docx")
        
        if DOC_AVAILABLE:
            formats.append(".doc")
        
        return formats
    
    def extract_text(self, file_path: Path) -> str:
        """Extract text from a document file."""
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        suffix = file_path.suffix.lower()
        
        if suffix not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {suffix}. Supported: {self.supported_formats}")
        
        text = ""
        
        try:
            if suffix == ".pdf":
                text = self._extract_pdf_text(file_path)
            elif suffix == ".docx":
                text = self._extract_docx_text(file_path)
            elif suffix == ".doc":
                text = self._extract_doc_text(file_path)
            elif suffix == ".txt":
                text = self._extract_txt_text(file_path)
            else:
                raise ValueError(f"No handler for file format: {suffix}")
        
        except Exception as e:
            console.error(f"Failed to extract text from {file_path.name}: {e}")
            return ""
        
        if self.clean_text:
            text = self._clean_extracted_text(text)
        
        return text
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF file using PyMuPDF."""
        if not PYMUPDF_AVAILABLE:
            raise ImportError("PyMuPDF not available for PDF processing")
        
        text_parts = []
        
        try:
            doc = fitz.open(file_path)
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Try text extraction first
                page_text = page.get_text()
                
                # If OCR is enabled and we got little/no text, try OCR
                if self.ocr_enabled and len(page_text.strip()) < 50:
                    try:
                        # Use OCR on the page
                        textpage = fitz.utils.get_textpage_ocr(page)
                        if textpage:
                            page_text = textpage.extractText()
                    except Exception as e:
                        console.log(f"OCR failed for page {page_num + 1}: {e}")
                
                if page_text.strip():
                    text_parts.append(page_text)
            
            doc.close()
            
        except Exception as e:
            raise Exception(f"PDF processing failed: {e}")
        
        return "\n\n".join(text_parts)
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX file using python-docx."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available for DOCX processing")
        
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            raise Exception(f"DOCX processing failed: {e}")
    
    def _extract_doc_text(self, file_path: Path) -> str:
        """Extract text from DOC file using docx2txt."""
        if not DOC_AVAILABLE:
            raise ImportError("docx2txt not available for DOC processing")
        
        try:
            text = docx2txt.process(str(file_path))
            return text or ""
        except Exception as e:
            raise Exception(f"DOC processing failed: {e}")
    
    def _extract_txt_text(self, file_path: Path) -> str:
        """Extract text from plain text file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, try with error handling
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                return f.read()
                
        except Exception as e:
            raise Exception(f"TXT processing failed: {e}")
    
    def _clean_extracted_text(self, text: str) -> str:
        """Clean extracted text by removing artifacts and normalizing."""
        if not text:
            return text
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove common PDF artifacts
        pattern = r'[^\w\s\.\,\!\?\;\:\-\(\)\[\]\"\'\/]'
        text = re.sub(pattern, ' ', text)
        
        # Remove repeated characters (common in OCR artifacts)
        text = re.sub(r'(.)\1{4,}', r'\1\1', text)
        
        # Remove very short lines (often artifacts)
        lines = text.split('\n')
        cleaned_lines = [li for li in lines if (li := line.strip()) and (len(li) > 3)]
        
        # Join lines back together
        text = '\n'.join(cleaned_lines)
        
        # Normalize spacing
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = text.strip()
        
        return text
    
    def get_file_info(self, file_path: Path) -> Dict[str, Any]:
        """Get information about a file."""
        if not file_path.exists():
            return {}
        
        info = {
            "name": file_path.name,
            "size": file_path.stat().st_size,
            "format": file_path.suffix.lower(),
            "supported": file_path.suffix.lower() in self.supported_formats
        }
        
        # Add format-specific info
        if file_path.suffix.lower() == ".pdf" and PYMUPDF_AVAILABLE:
            try:
                doc = fitz.open(file_path)
                info["pages"] = len(doc)
                doc.close()
            except:
                pass
        
        return info
    
    def batch_extract(self, file_paths: List[Path]) -> Dict[Path, str]:
        """Extract text from multiple files."""
        results = {}
        
        for file_path in file_paths:
            try:
                text = self.extract_text(file_path)
                results[file_path] = text
            except Exception as e:
                console.error(f"Failed to process {file_path.name}: {e}")
                results[file_path] = ""
        
        return results
